from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from .types import RawDocument


def _clean_text(text: str) -> str:
    text = text.replace("\x00", " ").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _display_path(path: Path) -> str:
    try:
        return str(path.relative_to(Path.cwd()))
    except ValueError:
        return path.name


class DirectoryDocumentLoader:
    def __init__(self, supported_extensions: tuple[str, ...]) -> None:
        self.supported_extensions = tuple(ext.lower() for ext in supported_extensions)

    def load_directory(self, directory: Path) -> list[RawDocument]:
        if not directory.exists():
            raise FileNotFoundError(f"Data directory does not exist: {directory}")

        files = sorted(
            path for path in directory.rglob("*") if path.is_file() and path.suffix.lower() in self.supported_extensions
        )
        if not files:
            raise FileNotFoundError(
                f"Khong tim thay file .json/.jsonl/.pdf/.docx trong thu muc {directory}"
            )

        documents: list[RawDocument] = []
        for path in files:
            documents.extend(self._load_file(path))

        if not documents:
            raise ValueError(f"Khong extract duoc noi dung tu cac file trong {directory}")
        return documents

    def _load_file(self, path: Path) -> list[RawDocument]:
        suffix = path.suffix.lower()
        if suffix == ".json":
            return self._load_json(path)
        if suffix == ".jsonl":
            return self._load_jsonl(path)
        if suffix == ".pdf":
            return self._load_pdf(path)
        if suffix == ".docx":
            return self._load_docx(path)
        return []

    def _load_json(self, path: Path) -> list[RawDocument]:
        payload = json.loads(path.read_text(encoding="utf-8"))
        text = self._build_json_text(
            header_lines=[f"Tep JSON: {path.name}"],
            payload=payload,
        )
        if not text:
            return []
        return [
            RawDocument(
                source_id=path.stem,
                source_name=path.name,
                source_type="json",
                text=text,
                metadata={"source_path": _display_path(path)},
            )
        ]

    def _load_jsonl(self, path: Path) -> list[RawDocument]:
        documents: list[RawDocument] = []
        with path.open("r", encoding="utf-8") as handle:
            for line_number, raw_line in enumerate(handle, start=1):
                line = raw_line.strip()
                if not line:
                    continue
                payloads = self._parse_jsonl_payloads(path=path, line=line, line_number=line_number)
                for record_index, payload in enumerate(payloads, start=1):
                    if not isinstance(payload, dict):
                        payload = {"content": payload}

                    record_id = str(payload.get("id") or f"{path.stem}-line-{line_number}-{record_index}")
                    source_id = f"{path.stem}-{line_number}-{record_index}-{record_id}"
                    header_lines = [
                        f"Tep JSONL: {path.name}",
                        f"Dong: {line_number}",
                    ]
                    if len(payloads) > 1:
                        header_lines.append(f"Ban ghi tren dong: {record_index}")

                    text = self._build_json_text(header_lines=header_lines, payload=payload)
                    if not text:
                        continue

                    documents.append(
                        RawDocument(
                            source_id=source_id,
                            source_name=path.name,
                            source_type="jsonl",
                            text=text,
                            metadata={
                                "source_path": _display_path(path),
                                "line_number": line_number,
                                "record_index": record_index,
                                "record_id": record_id,
                                "record_title": str(payload.get("title", "")).strip(),
                                "record_category": str(payload.get("category", "")).strip(),
                                "record_source": str(payload.get("source", "")).strip(),
                            },
                        )
                    )
        return documents

    def _load_pdf(self, path: Path) -> list[RawDocument]:
        reader = PdfReader(str(path))
        documents: list[RawDocument] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = _clean_text(page.extract_text() or "")
            if not text:
                continue
            documents.append(
                RawDocument(
                    source_id=f"{path.stem}-page-{page_number}",
                    source_name=path.name,
                    source_type="pdf",
                    text=text,
                    metadata={
                        "page_number": page_number,
                        "source_path": _display_path(path),
                    },
                )
            )
        return documents

    def _load_docx(self, path: Path) -> list[RawDocument]:
        document = DocxDocument(str(path))
        paragraphs = [_clean_text(paragraph.text) for paragraph in document.paragraphs if paragraph.text.strip()]

        table_lines: list[str] = []
        for table in document.tables:
            for row in table.rows:
                values = [_clean_text(cell.text) for cell in row.cells if cell.text.strip()]
                if values:
                    table_lines.append(" | ".join(values))

        text = _clean_text("\n".join(paragraphs + ([""] if table_lines else []) + table_lines))
        if not text:
            return []
        return [
            RawDocument(
                source_id=path.stem,
                source_name=path.name,
                source_type="docx",
                text=text,
                metadata={"source_path": _display_path(path)},
            )
        ]

    def _render_json_lines(self, value: object, depth: int = 0) -> list[str]:
        lines: list[str] = []

        if isinstance(value, dict):
            for key, child in value.items():
                label = str(key).replace("_", " ")
                if isinstance(child, (dict, list)):
                    prefix = "#" * min(depth + 1, 6)
                    lines.append(f"{prefix} {label}")
                    lines.extend(self._render_json_lines(child, depth + 1))
                else:
                    lines.append(f"- {label}: {self._format_scalar(child)}")
            return lines

        if isinstance(value, list):
            for index, item in enumerate(value, start=1):
                if isinstance(item, (dict, list)):
                    lines.append(f"- muc {index}")
                    nested = self._render_json_lines(item, depth + 1)
                    lines.extend(f"  {line}" for line in nested)
                else:
                    lines.append(f"- {self._format_scalar(item)}")
            return lines

        return [self._format_scalar(value)]

    @staticmethod
    def _format_scalar(value: object) -> str:
        if value is None:
            return ""
        if isinstance(value, bool):
            return "true" if value else "false"
        return str(value).strip()

    def _build_json_text(self, *, header_lines: list[str], payload: object) -> str:
        lines = [*header_lines, ""]
        lines.extend(self._render_json_lines(payload))
        return _clean_text("\n".join(lines))

    def _parse_jsonl_payloads(self, *, path: Path, line: str, line_number: int) -> list[object]:
        decoder = json.JSONDecoder()
        payloads: list[object] = []
        position = 0
        line_length = len(line)

        while position < line_length:
            while position < line_length and line[position].isspace():
                position += 1
            if position >= line_length:
                break

            try:
                payload, end_position = decoder.raw_decode(line, position)
            except json.JSONDecodeError as exc:
                if payloads:
                    # Ignore trailing junk after at least one valid JSON object.
                    break
                raise ValueError(
                    f"JSONL khong hop le tai {path.name}, dong {line_number}: {exc}"
                ) from exc

            payloads.append(payload)
            position = end_position

        if not payloads:
            raise ValueError(
                f"JSONL khong hop le tai {path.name}, dong {line_number}: khong doc duoc object nao"
            )
        return payloads
